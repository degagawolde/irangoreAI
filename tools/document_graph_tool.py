"""Document loading, chunking, and graph ingestion tools for GraphRAG."""

from __future__ import annotations

import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import time

from config import get_settings
from core.exceptions import GraphException
from core.logger import get_logger
from graph import get_graph
from llms import get_embeddings

from pypdf import PdfReader
from docx import Document

from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = get_logger(__name__)


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".rst",".docx",".pdf"}
PDF_PAGE_BREAK = "\n\n<<<PAGE_BREAK>>>\n\n"


class DocumentGraphIngestionTool:
    """Tool to load documents, chunk them, and ingest into Neo4j graph."""

    def __init__(self) -> None:
        self.settings = get_settings()
        self.graph = get_graph()
        self.embeddings = get_embeddings()

    def load_documents(self, path: str) -> List[Dict[str, Any]]:
        """Load text-like documents from a file or directory."""
        root = Path(path)
        if not root.exists():
            raise GraphException(f"Path does not exist: {path}")

        files: List[Path] = []
        if root.is_file():
            files = [root]
        else:
            files = [
                p
                for p in root.rglob("*")
                if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
            ]

        documents: List[Dict[str, Any]] = []
        for file_path in files:
            try:
                payload = self._read_document_payload(file_path)
                text = payload["content"].strip()
                if not text:
                    logger.warning(f"Skipping empty document: {file_path}")
                    continue
                char_count = len(text)
                logger.info(
                    "Loaded document: title=%s path=%s chars=%s",
                    file_path.name,
                    str(file_path.resolve()),
                    char_count,
                )
                documents.append(
                    {
                        "source_path": str(file_path.resolve()),
                        "title": file_path.name,
                        "content": text,
                        "page_spans": payload.get("page_spans", []),
                    }
                )
            except Exception as e:
                logger.warning(f"Skipping unreadable file {file_path}: {str(e)}")

        if not documents:
            raise GraphException(
                f"No supported documents found in {path}. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
            )

        logger.info(f"Loaded {len(documents)} documents from {path}")
        return documents

    def chunk_documents(
        self,
        documents: List[Dict[str, Any]],
        chunk_size: int = 1000,
        chunk_overlap: int = 150,
    ) -> List[Dict[str, Any]]:
        """Chunk loaded documents for graph/vector storage."""
        if chunk_overlap >= chunk_size:
            raise GraphException("chunk_overlap must be smaller than chunk_size")

        try:
           
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", ". ", " ", ""],
            )
        except Exception as e:
            raise GraphException(f"Failed to initialize text splitter: {str(e)}")

        chunk_rows: List[Dict[str, Any]] = []
        for doc in documents:
            document_id = hashlib.sha256(doc["source_path"].encode("utf-8")).hexdigest()
            chunks = splitter.split_text(doc["content"])
            chunk_spans = self._build_chunk_spans(doc["content"], chunks)
            logger.info(
                "Chunking document: title=%s doc_id=%s total_chunks=%s",
                doc["title"],
                document_id,
                len(chunks),
            )
            for idx, (chunk_text, span) in enumerate(zip(chunks, chunk_spans)):
                normalized = chunk_text.strip()
                if not normalized:
                    logger.debug(
                        "Skipping empty chunk: title=%s doc_id=%s chunk_index=%s",
                        doc["title"],
                        document_id,
                        idx,
                    )
                    continue
                logger.info(
                    "Prepared chunk: title=%s doc_id=%s chunk_id=%s chunk_index=%s chars=%s",
                    doc["title"],
                    document_id,
                    f"{document_id}:{idx}",
                    idx,
                    len(normalized),
                )
                chunk_rows.append(
                    {
                        "document_id": document_id,
                        "document_title": doc["title"],
                        "source_path": doc["source_path"],
                        "chunk_index": idx,
                        "chunk_id": f"{document_id}:{idx}",
                        "text": normalized,
                        "page_number": self._resolve_page_number(
                            doc.get("page_spans", []), span[0]
                        ),
                        "line_number": self._line_number_for_offset(doc["content"], span[0]),
                    }
                )

        if not chunk_rows:
            raise GraphException("Chunking produced no content")

        logger.info(f"Created {len(chunk_rows)} chunks from {len(documents)} documents")
        return chunk_rows

    def ingest_chunks(self, chunks: List[Dict[str, Any]]) -> Dict[str, int]:
        """Ingest chunks into Neo4j as a document knowledge graph."""
        created_at = datetime.now(timezone.utc).isoformat()

        texts = [c["text"] for c in chunks]
        vectors = self._embed_in_batches(texts)
        if len(vectors) != len(chunks):
            raise GraphException("Embedding count did not match chunk count")

        rows: List[Dict[str, Any]] = []
        for chunk, embedding in zip(chunks, vectors):
            logger.info(
                "Embedding chunk: chunk_id=%s doc_id=%s chunk_index=%s embedding_dimensions=%s",
                chunk["chunk_id"],
                chunk["document_id"],
                chunk["chunk_index"],
                len(embedding),
            )
            row = {**chunk, "embedding": embedding, "created_at": created_at}
            rows.append(row)

        self._ensure_vector_index(dimensions=len(vectors[0]))

        upsert_query = """
UNWIND $rows AS row
MERGE (d:Document {id: row.document_id})
  ON CREATE SET
    d.title = row.document_title,
    d.source_path = row.source_path,
    d.created_at = row.created_at
  ON MATCH SET
    d.title = row.document_title,
    d.source_path = row.source_path
MERGE (c:Chunk {id: row.chunk_id})
  ON CREATE SET
    c.created_at = row.created_at
SET
  c.document_id = row.document_id,
  c.chunk_index = row.chunk_index,
  c.text = row.text,
  c.embedding = row.embedding,
  c.source_path = row.source_path,
  c.page_number = row.page_number,
  c.line_number = row.line_number
MERGE (d)-[:HAS_CHUNK]->(c)
WITH row, c
OPTIONAL MATCH (prev:Chunk {id: row.document_id + ':' + toString(row.chunk_index - 1)})
FOREACH (_ IN CASE WHEN prev IS NULL THEN [] ELSE [1] END |
  MERGE (prev)-[:NEXT_CHUNK]->(c)
)
"""
        self.graph.query(upsert_query, params={"rows": rows})
        for row in rows:
            logger.info(
                "Ingested chunk to graph: chunk_id=%s doc_id=%s chunk_index=%s source_path=%s",
                row["chunk_id"],
                row["document_id"],
                row["chunk_index"],
                row["source_path"],
            )

        doc_count = len({c["document_id"] for c in chunks})
        chunk_count = len(chunks)
        logger.info(f"Ingested {doc_count} documents and {chunk_count} chunks")
        return {"documents": doc_count, "chunks": chunk_count}

    def ingest_documents(
        self,
        path: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 150,
    ) -> Dict[str, int]:
        """Full ingestion flow: load -> chunk -> embed -> ingest."""
        documents = self.load_documents(path)
        chunks = self.chunk_documents(
            documents=documents,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        return self.ingest_chunks(chunks)

    def _ensure_vector_index(self, dimensions: int) -> None:
        """Ensure vector index exists for chunk embeddings."""
        query = f"""
CREATE VECTOR INDEX {self.settings.VECTOR_INDEX_NAME} IF NOT EXISTS
FOR (c:{self.settings.VECTOR_NODE_LABEL}) ON (c.{self.settings.VECTOR_EMBEDDING_PROPERTY})
OPTIONS {{
  indexConfig: {{
    `vector.dimensions`: $dimensions,
    `vector.similarity_function`: 'cosine'
  }}
}}
"""
        try:
            self.graph.query(query, params={"dimensions": dimensions})
        except Exception as e:
            raise GraphException(f"Failed to create/verify vector index: {str(e)}")

    def _read_document_payload(self, file_path: Path) -> Dict[str, Any]:
        """Read text + optional page spans from supported files."""
        suffix = file_path.suffix.lower()
        if suffix in {".txt", ".md", ".markdown", ".rst"}:
            return {"content": file_path.read_text(encoding="utf-8", errors="ignore"), "page_spans": []}

        if suffix == ".pdf":
            reader = PdfReader(str(file_path))
            page_texts = [(page.extract_text() or "") for page in reader.pages]
            extracted = PDF_PAGE_BREAK.join(page_texts).strip()
            if extracted:
                cleaned_content, page_spans = self._build_pdf_page_spans_from_joined_text(extracted)
                return {"content": cleaned_content, "page_spans": page_spans}

            if self.settings.PDF_OCR_FALLBACK_ENABLED:
                logger.warning(
                    "No extractable text found in PDF. Falling back to OCR: %s",
                    str(file_path),
                )
                ocr_text = self._extract_pdf_text_with_ocr(file_path)
                if ocr_text.strip():
                    cleaned_content, page_spans = self._build_pdf_page_spans_from_joined_text(ocr_text)
                    return {"content": cleaned_content, "page_spans": page_spans}
                logger.warning("OCR fallback returned empty text for PDF: %s", str(file_path))
            return {"content": extracted, "page_spans": []}

        if suffix == ".docx":
            doc = Document(str(file_path))
            return {"content": "\n".join(p.text for p in doc.paragraphs if p.text), "page_spans": []}

        raise GraphException(f"Unsupported extension encountered: {suffix}")

    def _extract_pdf_text_with_ocr(self, file_path: Path) -> str:
        """OCR fallback for scanned PDFs when native text extraction fails."""
        try:
            import pytesseract
            import pypdfium2 as pdfium
        except ImportError as e:
            logger.warning(
                "OCR fallback unavailable for %s due to missing packages: %s",
                str(file_path),
                str(e),
            )
            return ""

        try:
            pdf = pdfium.PdfDocument(str(file_path))
            ocr_chunks: List[str] = []
            for page_index in range(len(pdf)):
                page = pdf[page_index]
                bitmap = page.render(scale=2.0).to_pil()
                text = pytesseract.image_to_string(bitmap, lang=self.settings.PDF_OCR_LANG) or ""
                cleaned = text.strip()
                if cleaned:
                    ocr_chunks.append(cleaned)
            return PDF_PAGE_BREAK.join(ocr_chunks)
        except Exception as e:
            logger.warning("OCR fallback failed for %s: %s", str(file_path), str(e))
            return ""

    def _build_chunk_spans(self, content: str, chunks: List[str]) -> List[Tuple[int, int]]:
        """Map each chunk to (start, end) offsets in original content."""
        spans: List[Tuple[int, int]] = []
        cursor = 0
        for chunk in chunks:
            if not chunk:
                spans.append((cursor, cursor))
                continue
            idx = content.find(chunk, cursor)
            if idx == -1:
                idx = content.find(chunk)
            if idx == -1:
                idx = cursor
            end = idx + len(chunk)
            spans.append((idx, end))
            cursor = max(cursor, end - 1)
        return spans

    def _line_number_for_offset(self, content: str, offset: int) -> int:
        """Convert character offset to 1-based line number."""
        if offset <= 0:
            return 1
        return content.count("\n", 0, min(offset, len(content))) + 1

    def _build_pdf_page_spans_from_joined_text(self, joined_text: str) -> Tuple[str, List[Dict[str, int]]]:
        """Create page spans from marker-delimited PDF text."""
        parts = joined_text.split(PDF_PAGE_BREAK)
        content_parts: List[str] = []
        spans: List[Dict[str, int]] = []
        cursor = 0
        for page_idx, part in enumerate(parts, start=1):
            clean = part.strip()
            if not clean:
                continue
            start = cursor
            end = start + len(clean)
            spans.append({"page_number": page_idx, "start": start, "end": end})
            content_parts.append(clean)
            cursor = end + 1
        return "\n".join(content_parts), spans

    def _resolve_page_number(self, page_spans: List[Dict[str, int]], offset: int) -> Optional[int]:
        """Resolve page number from char offset."""
        for span in page_spans:
            if span["start"] <= offset <= span["end"]:
                return span["page_number"]
        return None

    def _embed_in_batches(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings in batches to avoid oversized requests/hangs."""
        batch_size = max(1, int(self.settings.EMBEDDING_BATCH_SIZE))
        total = len(texts)
        logger.info(
            "Starting embeddings: total_texts=%s batch_size=%s model=%s",
            total,
            batch_size,
            self.settings.EMBEDDING_MODEL,
        )

        vectors: List[List[float]] = []
        started_at = time.perf_counter()
        for start in range(0, total, batch_size):
            end = min(start + batch_size, total)
            batch = texts[start:end]
            batch_started_at = time.perf_counter()
            logger.info("Embedding batch started: start=%s end=%s size=%s", start, end - 1, len(batch))
            try:
                batch_vectors = self.embeddings.embed_documents(batch)
            except Exception as e:
                raise GraphException(
                    f"Embedding batch failed for indexes {start}-{end - 1}: {str(e)}"
                ) from e

            if len(batch_vectors) != len(batch):
                raise GraphException(
                    f"Embedding batch size mismatch for indexes {start}-{end - 1}: "
                    f"expected={len(batch)} got={len(batch_vectors)}"
                )

            vectors.extend(batch_vectors)
            logger.info(
                "Embedding batch finished: start=%s end=%s elapsed_seconds=%.2f",
                start,
                end - 1,
                time.perf_counter() - batch_started_at,
            )

        logger.info(
            "Completed embeddings: total_vectors=%s elapsed_seconds=%.2f",
            len(vectors),
            time.perf_counter() - started_at,
        )
        return vectors


def get_document_graph_ingestion_tool() -> DocumentGraphIngestionTool:
    """Get a document graph ingestion tool instance."""
    return DocumentGraphIngestionTool()


def ingest_documents_to_graph(
    path: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 150,
) -> Dict[str, int]:
    """Convenience function to ingest documents from filesystem into graph."""
    tool = get_document_graph_ingestion_tool()
    return tool.ingest_documents(path, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def ingest_documents_tool(payload: str) -> str:
    """Agent-friendly wrapper for document ingestion tool.

    Expected payload JSON:
    {"path": "./docs", "chunk_size": 1000, "chunk_overlap": 150}
    """
    try:
        data = json.loads(payload)
        path = data["path"]
        chunk_size = int(data.get("chunk_size", 1000))
        chunk_overlap = int(data.get("chunk_overlap", 150))

        result = ingest_documents_to_graph(
            path=path,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        return (
            f"Ingestion complete. Documents: {result['documents']}, "
            f"Chunks: {result['chunks']}"
        )
    except KeyError:
        return "Invalid payload. Provide JSON with at least: {'path': '...'}"
    except Exception as e:
        return f"Document ingestion failed: {str(e)}"
